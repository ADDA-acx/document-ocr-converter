from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_run_fonts(run, east_asia: str = "SimSun", ascii_font: str = "Arial") -> None:
    run.font.name = ascii_font
    run.font.size = Pt(11)
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)


def _format_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(11)


def _add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    _set_run_fonts(run)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            _set_run_fonts(run)


def write_ocr_docx(
    pages: list[dict[str, Any]],
    output_path: Path,
    keep_page_breaks: bool = True,
    append_page_images: bool = False,
) -> None:
    document = Document()
    _format_document(document)

    for page_index, page in enumerate(pages):
        if page_index > 0 and keep_page_breaks:
            document.add_page_break()
        for block in page.get("blocks", []):
            if block.get("type") == "table":
                _add_table(document, block.get("rows", []))
            else:
                text = str(block.get("text", "")).strip()
                if text:
                    _add_paragraph(document, text)

    if append_page_images:
        document.add_section(WD_SECTION.NEW_PAGE)
        _add_paragraph(document, "原页面截图（用于核对）")
        for page_index, page in enumerate(pages):
            image_path = page.get("image_path")
            if image_path and Path(image_path).exists():
                _add_paragraph(document, f"第 {page_index + 1} 页")
                document.add_picture(str(image_path), width=Cm(18))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
